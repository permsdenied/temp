import copy
import logging
import re

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{[\w_]+\}\}")

# xml:space attribute in Clark notation
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _replace_text_in_run(run, replacements: dict) -> None:
    """Replace placeholder tokens inside a single run, preserving formatting."""
    for key, value in replacements.items():
        token = "{{" + key + "}}"
        if token in run.text:
            new_text = run.text.replace(token, str(value))
            run.text = new_text
            wt = run._r.find(qn("w:t"))
            if wt is not None and new_text and (
                new_text[0] == " " or new_text[-1] == " "
            ):
                wt.set(_XML_SPACE, "preserve")


def _replace_in_paragraph(paragraph, replacements: dict) -> None:
    """Replace placeholders across all w:t elements of a paragraph.

    Works at the XML level so that:
    - Text inside hyperlinks, smart tags, and other wrapper elements is captured
      (not only direct-child runs returned by ``paragraph.runs``).
    - Placeholders split across multiple runs are handled correctly.
    - Formatting of runs that are *not* part of a placeholder is fully preserved
      because only the affected ``w:t`` elements are modified.

    Algorithm:
    1. Collect all ``w:t`` elements from the paragraph XML.
    2. Build a combined string and record the start boundary of each element.
    3. Find every ``{{key}}`` occurrence and record (start, end, value) intervals.
    4. For each ``w:t``, compute new text by:
       - Keeping unchanged characters verbatim.
       - Emitting the replacement value in the *first* ``w:t`` that contains the
         placeholder's start position; ``w:t`` elements that are entirely inside
         the matched range but do not own the start are cleared.
    5. Set ``xml:space="preserve"`` on any ``w:t`` whose text begins or ends with
       a space, so Word does not silently strip leading/trailing whitespace.
    """
    p_xml = paragraph._p
    wt_elements = p_xml.findall(".//" + qn("w:t"))

    if not wt_elements:
        return

    # Build combined text and element start-boundaries
    texts = [wt.text or "" for wt in wt_elements]
    combined_text = "".join(texts)

    if "{{" not in combined_text:
        return

    boundaries = []
    pos = 0
    for t in texts:
        boundaries.append(pos)
        pos += len(t)
    boundaries.append(pos)  # sentinel: total length

    # Collect all replacement intervals (start, end, new_value)
    intervals = []
    for key, value in replacements.items():
        token = "{{" + key + "}}"
        search_start = 0
        while True:
            idx = combined_text.find(token, search_start)
            if idx == -1:
                break
            intervals.append((idx, idx + len(token), str(value)))
            search_start = idx + len(token)

    if not intervals:
        return

    # Sort by start position and remove overlapping intervals
    intervals.sort(key=lambda x: x[0])
    resolved_intervals: list = []
    last_end = -1
    for start, end, value in intervals:
        if start >= last_end:
            resolved_intervals.append((start, end, value))
            last_end = end

    # Apply replacements to each w:t element independently
    for k, wt in enumerate(wt_elements):
        wt_start = boundaries[k]
        wt_end = boundaries[k + 1]

        if wt_start == wt_end:
            continue  # empty element — nothing to do

        element_text_parts: list = []
        cur = wt_start  # current read position in combined_text

        for rep_start, rep_end, rep_value in resolved_intervals:
            if rep_end <= wt_start:
                continue  # replacement is entirely before this element
            if rep_start >= wt_end:
                break  # replacement is entirely after this element

            # Emit any unchanged text that precedes this replacement
            if cur < rep_start:
                element_text_parts.append(combined_text[cur:rep_start])

            # Emit the replacement value only in the element that owns rep_start
            if rep_start >= wt_start:
                element_text_parts.append(rep_value)

            cur = rep_end  # advance past the replacement

        # Emit any remaining unchanged text inside this element
        if cur < wt_end:
            element_text_parts.append(combined_text[cur:wt_end])

        new_text = "".join(element_text_parts)

        if new_text != (wt.text or ""):
            wt.text = new_text

        # Preserve leading/trailing spaces (Word strips them by default)
        if new_text and (new_text[0] == " " or new_text[-1] == " "):
            wt.set(_XML_SPACE, "preserve")


def _iter_paragraphs(doc: Document):
    """Yield all paragraphs in the document including those inside tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _find_price_table(doc: Document):
    """Return the first table that contains a {{price_table}} marker, or None."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{price_table}}" in cell.text:
                    return table
    return None


def _fill_price_table(table, price_rows: list) -> None:
    """Fill the price table with data rows.

    Assumes the table has:
      - A header row (row 0) — left untouched.
      - A template data row (row 1, containing {{price_table}}) — used as the
        style template and then replaced.
      - Optionally a totals/footer row at the end — left untouched.

    Args:
        table: The docx Table object.
        price_rows: List of dicts with keys: item, qty, unit, price, total.
    """
    if len(table.rows) < 2:
        logger.warning("Price table has fewer than 2 rows — skipping fill.")
        return

    template_row = table.rows[1]
    # Index of the last row (totals footer) — keep it, insert before it
    # If the table only has header + template, there is no footer row.
    has_footer = len(table.rows) > 2

    # Remove the placeholder template row's cell texts so we can reuse its XML
    # as a style reference.
    template_row_xml = copy.deepcopy(template_row._tr)

    # Remove all existing data rows (keep header at index 0 and footer if any)
    rows_to_remove = list(table.rows[1:] if not has_footer else table.rows[1:-1])
    for row in rows_to_remove:
        table._tbl.remove(row._tr)

    footer_tr = table.rows[-1]._tr if has_footer else None

    for idx, price_row in enumerate(price_rows, start=1):
        new_tr = copy.deepcopy(template_row_xml)
        cells_xml = new_tr.findall(qn("w:tc"))
        # Expected columns: №, item, qty, unit, price, total  (6 columns)
        # Or: №, item, price  etc. — we map what we have.
        col_values = [
            str(idx),
            price_row.get("item", ""),
            price_row.get("qty", ""),
            price_row.get("unit", ""),
            price_row.get("price", ""),
            price_row.get("total", ""),
        ]
        for col_idx, cell_xml in enumerate(cells_xml):
            if col_idx >= len(col_values):
                break
            # Find the first run inside the cell and set its text
            runs = cell_xml.findall(".//" + qn("w:r"))
            t_elements = cell_xml.findall(".//" + qn("w:t"))
            if t_elements:
                # Clear all text elements, set first to value
                val = col_values[col_idx]
                t_elements[0].text = val
                if val and (val[0] == " " or val[-1] == " "):
                    t_elements[0].set(_XML_SPACE, "preserve")
                for t in t_elements[1:]:
                    t.text = ""
            elif runs:
                # Create a w:t element in the first run
                val = col_values[col_idx]
                t = etree.SubElement(runs[0], qn("w:t"))
                t.text = val
                if val and (val[0] == " " or val[-1] == " "):
                    t.set(_XML_SPACE, "preserve")

        # Insert before footer row if it exists, otherwise append
        tbl = table._tbl
        if footer_tr is not None:
            tbl.insert(list(tbl).index(footer_tr), new_tr)
        else:
            tbl.append(new_tr)


def fill_template(template_path: str, content: dict, output_path: str) -> None:
    """Fill a .docx template with KP content and save to output_path.

    Placeholders in the template use the ``{{key}}`` syntax.
    Images, headers, footers, and formatting are preserved.

    Args:
        template_path: Path to the .docx template file.
        content: Dict with KP content fields (see gemini_client.py for schema).
        output_path: Where to save the filled document.
    """
    doc = Document(template_path)

    # Simple scalar replacements (everything except price_table)
    scalar_replacements = {
        k: v for k, v in content.items() if k != "price_table"
    }

    # Replace in body paragraphs and table cells
    for paragraph in _iter_paragraphs(doc):
        _replace_in_paragraph(paragraph, scalar_replacements)

    # Replace in headers and footers (preserving images and formatting)
    for section in doc.sections:
        for header_footer in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, scalar_replacements)

    # Fill price table if a placeholder table exists
    price_rows = content.get("price_table", [])
    if price_rows:
        price_table = _find_price_table(doc)
        if price_table:
            _fill_price_table(price_table, price_rows)
        else:
            logger.warning(
                "No {{price_table}} marker found in document — "
                "price rows will not be inserted."
            )

    doc.save(output_path)
    logger.info("Filled template saved to %s", output_path)
